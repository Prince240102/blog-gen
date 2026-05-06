"""
API Routes — Auth, Chat (agentic streaming), Publish, Sessions.

The chat endpoint runs a ReAct agent that decides which tools to call.
Conversation history is rebuilt from saved user/assistant text pairs
each turn, with the current blog draft injected as context.
"""

from __future__ import annotations

import asyncio
import io
import json
import logging
import re

from typing import Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.security import OAuth2PasswordRequestForm
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from sse_starlette.sse import EventSourceResponse
from fastapi.responses import Response

from app.agents.agent import get_agent
from app.agents.publisher import run_publisher
from app.agents.tools import TOOL_DISPLAY
from app.core.auth import (
    create_access_token,
    get_current_user,
)
from app.models.schemas import (
    ChatRequest,
    ChatResponse,
    SessionResponse,
    Token,
    User,
    UserCreate,
)
from app.services import store
from app.services.llm import llm_fast
from app.services.progress import ProgressCallback, set_progress, get_progress
from fastapi import UploadFile, File as FileType

logger = logging.getLogger(__name__)
router = APIRouter()


def _tool_label(name: str) -> tuple[str, str]:
    return TOOL_DISPLAY.get(name, (name, "⚙️"))


def _get_owned_session(session_id: str, user_id: str):
    session = store.get_session(session_id)
    if not session:
        raise HTTPException(404, "Session not found")
    if session.user_id != user_id:
        raise HTTPException(403)
    return session


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------


@router.post("/auth/register", response_model=User)
async def register(user: UserCreate):
    existing = store.get_user_by_email(user.email)
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    return store.create_user(user.email, user.username, user.password)


@router.post("/auth/login", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = store.authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=401,
            detail="Incorrect email or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return Token(
        access_token=create_access_token(data={"sub": user["id"]}),
        token_type="bearer",
    )


# ---------------------------------------------------------------------------
# File upload
# ---------------------------------------------------------------------------


@router.post("/chat/upload")
async def upload_file(
    file: UploadFile = FileType(...),
    user_id: str = Depends(get_current_user),
):
    """Extract text from an uploaded PDF or DOCX file."""
    filename = file.filename or "unknown"
    suffix = filename.rsplit(".", 1)[-1].lower()

    if suffix not in ("pdf", "docx"):
        raise HTTPException(400, "Only .pdf and .docx files are supported")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:  # 10 MB limit
        raise HTTPException(400, "File too large (max 10 MB)")

    try:
        if suffix == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == "docx":
            import docx2txt
            text = docx2txt.process(io.BytesIO(content))
        else:
            text = ""
    except Exception as exc:
        raise HTTPException(400, f"Failed to extract text: {exc}")

    if not text.strip():
        raise HTTPException(400, "Could not extract any text from the file")

    return {"filename": filename, "text": text[:50000], "chars": len(text)}


# ---------------------------------------------------------------------------
# Chat (agentic streaming)
# ---------------------------------------------------------------------------

_CONTENT_TOOLS = {"write_blog", "humanize", "revise"}
# Only send blog_content event for these final tools, not intermediate ones
_FINAL_CONTENT_TOOLS = {"write_blog", "humanize", "revise"}
_DETAIL_TOOLS = {
    "web_search",
    "research_topic",
    "analyze_seo",
    "request_publish_approval",
    "publish_to_wordpress",
    "convert_to_vlog",
}
_CONTEXT_BUDGET = 50000
_RECENT_WINDOW = 5
_RECENT_CHAR_LIMIT = 8000
_DRAFT_CHAR_LIMIT = 15000


@router.post("/chat/stream")
async def chat_stream(request: ChatRequest, user_id: str = Depends(get_current_user)):
    sid = request.session_id
    if sid:
        _get_owned_session(sid, user_id)
    else:
        sid = store.create_session(user_id).session_id

    turn_id = store.add_message(sid, "user", request.message)
    input_messages = _build_input_messages(sid, request.message)
    agent = get_agent(sid)

    # Set up progress callback for this request
    progress = ProgressCallback()
    set_progress(progress)

    async def generate():
        agent_text = ""
        blog_title = ""
        blog_content = None
        needs_approval = False
        seo_keywords = ""
        seo_meta_desc = ""

        progress_queue = asyncio.Queue()
        step_seq = 0
        current_step_id: int | None = None

        async def poll_progress():
            while True:
                await asyncio.sleep(0.3)
                for msg in progress.drain():
                    await progress_queue.put(msg)

        poll_task = asyncio.create_task(poll_progress())

        try:
            async for event in agent.astream_events(
                {"messages": input_messages},
                version="v2",
            ):
                # Drain any pending progress messages first
                while not progress_queue.empty():
                    try:
                        msg = progress_queue.get_nowait()
                        if current_step_id is not None:
                            store.update_tool_step_progress(current_step_id, msg)
                        yield {
                            "event": "step_progress",
                            "data": json.dumps({"text": msg}),
                        }
                    except asyncio.QueueEmpty:
                        break

                kind = event["event"]
                metadata = event.get("metadata", {})
                node = metadata.get("langgraph_node", "")

                # --- Stream agent text tokens ---
                if kind == "on_chat_model_stream" and node == "agent":
                    chunk = event["data"]["chunk"]
                    if hasattr(chunk, "content") and chunk.content:
                        if isinstance(chunk.content, list):
                            text = "".join(
                                c.get("text", "") if isinstance(c, dict) else str(c)
                                for c in chunk.content
                            )
                        else:
                            text = str(chunk.content)
                        if text:
                            agent_text += text
                            yield {
                                "event": "token",
                                "data": json.dumps({"token": text}),
                            }

                # --- Tool started ---
                elif kind == "on_tool_start":
                    tool_name = event.get("name", "")
                    label, icon = _tool_label(tool_name)
                    step_seq += 1
                    current_step_id = store.start_tool_step(
                        sid,
                        turn_id=turn_id,
                        seq=step_seq,
                        tool=tool_name,
                        label=label,
                        icon=icon,
                    )
                    yield {
                        "event": "tool_start",
                        "data": json.dumps(
                            {"tool": tool_name, "label": label, "icon": icon}
                        ),
                    }

                # --- Tool errored ---
                elif kind == "on_tool_error":
                    tool_name = event.get("name", "")
                    err = event.get("data", {}).get("error")
                    err_str = str(err) if err is not None else "Tool error"
                    if current_step_id is not None:
                        store.finish_tool_step(current_step_id, status="error", output=err_str)
                        current_step_id = None
                    label, icon = _tool_label(tool_name)
                    yield {
                        "event": "tool_output",
                        "data": json.dumps({"tool": tool_name, "content": err_str}),
                    }
                    yield {
                        "event": "tool_end",
                        "data": json.dumps(
                            {"tool": tool_name, "label": label, "icon": icon, "status": "error"}
                        ),
                    }

                # --- Tool finished ---
                elif kind == "on_tool_end":
                    tool_name = event.get("name", "")
                    raw_output = event["data"]["output"]
                    output_str = (
                        raw_output.content
                        if hasattr(raw_output, "content")
                        else str(raw_output)
                    )

                    if tool_name == "analyze_seo":
                        try:
                            seo_data = json.loads(output_str)
                            seo_keywords = ",".join(seo_data.get("keywords", []))
                            seo_meta_desc = seo_data.get("meta_description", "")
                        except (json.JSONDecodeError, TypeError):
                            seo_keywords = ""
                            seo_meta_desc = ""

                    if tool_name in _CONTENT_TOOLS:
                        title_match = re.search(
                            r"^#\s+(.+)$", output_str, re.MULTILINE
                        )
                        if title_match:
                            blog_title = title_match.group(1).strip()
                        blog_content = output_str

                        # Save draft after first content tool (for subsequent tools in same turn)
                        # Don't create version yet - only at end of stream
                        if tool_name == "write_blog" and blog_content and sid:
                            store.set_draft(
                                sid,
                                title=blog_title,
                                content=blog_content,
                                word_count=len(blog_content.split()),
                                keywords=seo_keywords,
                                meta_description=seo_meta_desc,
                                create_version=False,
                            )

                        if tool_name in _FINAL_CONTENT_TOOLS:
                            yield {
                                "event": "blog_content",
                                "data": json.dumps({
                                    "title": blog_title,
                                    "content": blog_content,
                                    "word_count": len(blog_content.split()),
                                }),
                            }

                    if tool_name == "request_publish_approval":
                        needs_approval = True

                    if tool_name == "publish_to_wordpress":
                        if "Published as draft" in output_str or "Post ID" in output_str:
                            post_id = 0
                            permalink = ""
                            post_match = re.search(r"Post ID[:\s]+(\d+)", output_str)
                            link_match = re.search(r"(https?://\S+)", output_str)
                            if post_match:
                                post_id = int(post_match.group(1))
                            if link_match:
                                permalink = link_match.group(1).rstrip(".,)")
                            store.mark_published(sid, post_id, permalink)

                    label, icon = _tool_label(tool_name)

                    # Drain progress messages BEFORE completing the step
                    for msg in progress.drain():
                        if current_step_id is not None:
                            store.update_tool_step_progress(current_step_id, msg)
                        yield {
                            "event": "step_progress",
                            "data": json.dumps({"text": msg}),
                        }

                    if tool_name in _DETAIL_TOOLS and output_str:
                        yield {
                            "event": "tool_output",
                            "data": json.dumps(
                                {
                                    "tool": tool_name,
                                    "content": output_str[:8000],
                                }
                            ),
                        }

                    if current_step_id is not None:
                        # Heuristic: non-empty output means success.
                        store.finish_tool_step(
                            current_step_id,
                            status="error" if output_str.lstrip().startswith("Error") else "done",
                            output=output_str,
                        )
                        current_step_id = None

                    yield {
                        "event": "tool_end",
                        "data": json.dumps(
                            {"tool": tool_name, "label": label, "icon": icon}
                        ),
                    }

        finally:
            poll_task.cancel()
            if current_step_id is not None:
                store.finish_tool_step(current_step_id, status="aborted", output="aborted")
                current_step_id = None

        # Merge blog content and agent text into a single message
        if blog_content and agent_text.strip():
            store.add_message(sid, "assistant", f"{agent_text.strip()}\n\n---\n\n{blog_content[:15000]}")
        elif blog_content:
            store.add_message(sid, "assistant", blog_content[:15000])
        elif agent_text.strip():
            store.add_message(sid, "assistant", agent_text[:8000])

        if blog_content:
            store.set_draft(
                sid,
                title=blog_title,
                content=blog_content,
                word_count=len(blog_content.split()),
                keywords=seo_keywords,
                meta_description=seo_meta_desc,
            )

        versions = store.list_versions(sid) if blog_content else []

        yield {
            "event": "done",
            "data": json.dumps(
                {
                    "session_id": sid,
                    "can_publish": blog_content is not None or needs_approval,
                    "needs_approval": needs_approval,
                    "blog_title": blog_title if needs_approval else "",
                    "versions": [{"version": v["version"], "title": v["title"], "word_count": v["word_count"], "is_current": v["is_current"]} for v in versions],
                }
            ),
        }

    return EventSourceResponse(generate())


def _build_input_messages(session_id: str, current_message: str) -> list:
    """Build LangChain message list from session history with smart context.

    - Recent messages (last 5): full content, up to 8000 chars each
    - Older messages: summarized via llm_fast if no cached summary
    - Current draft: injected as SystemMessage (up to 15000 chars)
    - Total budget: ~50000 chars
    """
    messages = []
    history = store.get_recent_messages(session_id, limit=30)

    # Remove the message we just saved (it's also the current_message)
    if history and history[-1]["content"] == current_message:
        history = history[:-1]

    if len(history) <= _RECENT_WINDOW:
        # Short history — include everything
        for msg in history:
            content = msg["content"][:_RECENT_CHAR_LIMIT]
            if msg["role"] == "user":
                messages.append(HumanMessage(content=content))
            elif msg["role"] == "assistant":
                messages.append(AIMessage(content=content))
    else:
        # Split into older (summarized) and recent (full)
        older = history[:-_RECENT_WINDOW]
        recent = history[-_RECENT_WINDOW:]

        # Summarize older messages
        for msg in older:
            summary = msg.get("summary")
            if summary:
                text = f"[{msg['role'].title()}: {summary}]"
            elif msg["content"].strip():
                # Generate and cache summary
                summary = _summarize_message(msg["role"], msg["content"][:2000])
                if msg.get("id"):
                    store.update_message_summary(msg["id"], summary)
                text = f"[{msg['role'].title()}: {summary}]"
            else:
                continue
            if msg["role"] == "user":
                messages.append(HumanMessage(content=text))
            elif msg["role"] == "assistant":
                messages.append(AIMessage(content=text))

        # Recent messages in full
        for msg in recent:
            content = msg["content"][:_RECENT_CHAR_LIMIT]
            if msg["role"] == "user":
                messages.append(HumanMessage(content=content))
            elif msg["role"] == "assistant":
                messages.append(AIMessage(content=content))

    # Inject current blog draft as SystemMessage
    draft = store.get_draft(session_id)
    if draft and draft.get("humanized_content"):
        draft_title = draft.get("blog_title", "Untitled")
        draft_text = draft["humanized_content"][:_DRAFT_CHAR_LIMIT]
        draft_kw = draft.get("keywords", "")
        draft_meta = draft.get("meta_description", "")
        messages.append(
            SystemMessage(
                content=(
                    f"[Current session draft]\n"
                    f"Title: {draft_title}\n"
                    f"Keywords: {draft_kw}\n"
                    f"Meta description: {draft_meta}\n"
                    f"Word count: {draft.get('blog_word_count', 0)}\n\n"
                    f"{draft_text}"
                )
            )
        )

    messages.append(HumanMessage(content=current_message))
    return messages


def _summarize_message(role: str, content: str) -> str:
    """Generate a one-line summary of a message using llm_fast."""
    try:
        resp = llm_fast.invoke(
            [
                HumanMessage(
                    content=(
                        f"Summarize this {role} message in one short sentence (max 15 words). "
                        f"Just output the summary, nothing else:\n\n{content[:1500]}"
                    )
                )
            ]
        )
        return resp.content.strip()[:120]
    except Exception:
        # Fallback: truncate
        return content[:80].replace("\n", " ").strip() + "..."


# ---------------------------------------------------------------------------
# Chat (non-streaming fallback)
# ---------------------------------------------------------------------------


@router.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, user_id: str = Depends(get_current_user)):
    sid = request.session_id
    if sid:
        _get_owned_session(sid, user_id)
    else:
        sid = store.create_session(user_id).session_id
    store.add_message(sid, "user", request.message)

    input_messages = _build_input_messages(sid, request.message)
    agent = get_agent(sid)

    result = await asyncio.to_thread(
        agent.invoke, {"messages": input_messages}
    )
    updated = result.get("messages", input_messages)

    ai_text = ""
    for msg in reversed(updated):
        if isinstance(msg, AIMessage) and msg.content:
            ai_text = msg.content
            break

    store.add_message(sid, "assistant", ai_text[:8000])
    return ChatResponse(session_id=sid, message=ai_text, agent="agent")


# ---------------------------------------------------------------------------
# Publish (direct — for frontend button)
# ---------------------------------------------------------------------------


@router.post("/chat/publish")
async def publish_blog(request: ChatRequest, user_id: str = Depends(get_current_user)):
    sid = request.session_id
    if not sid:
        raise HTTPException(400, "session_id required")
    _get_owned_session(sid, user_id)
    draft = store.get_draft(sid)
    if not draft or not draft.get("humanized_content"):
        raise HTTPException(400, "No blog content in session")

    content = draft.get("humanized_content", draft.get("blog_content", ""))
    keywords = draft.get("keywords", "")
    meta_desc = draft.get("meta_description", "")

    meta = {}
    if draft.get("blog_title"):
        meta["rank_math_title"] = draft["blog_title"]
    if meta_desc:
        meta["rank_math_description"] = meta_desc
    if keywords:
        meta["rank_math_focus_keyword"] = keywords

    r = await asyncio.to_thread(
        run_publisher,
        title=draft.get("blog_title", "Untitled"),
        content=content,
        excerpt=meta_desc,
        status="draft",
        meta=meta if meta else None,
    )
    if r.get("success"):
        store.mark_published(sid, r.get("post_id", 0), r.get("permalink", ""))
    return {
        "success": r.get("success", False),
        "post_id": r.get("post_id"),
        "permalink": r.get("permalink"),
        "error": r.get("error"),
    }


# ---------------------------------------------------------------------------
# Versions
# ---------------------------------------------------------------------------


@router.get("/chat/versions/{session_id}")
async def get_versions(session_id: str, user_id: str = Depends(get_current_user)):
    s = store.get_session(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    if s.user_id != user_id:
        raise HTTPException(403)
    return store.list_versions(session_id)


@router.get("/chat/steps/{session_id}")
async def get_steps(session_id: str, user_id: str = Depends(get_current_user)):
    s = store.get_session(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    if s.user_id != user_id:
        raise HTTPException(403)
    return store.list_tool_steps(session_id, limit=80)


@router.get("/chat/version-content/{session_id}")
async def get_version_content(session_id: str, version: int, user_id: str = Depends(get_current_user)):
    s = store.get_session(session_id)
    if not s:
        raise HTTPException(404, "Session not found")
    if s.user_id != user_id:
        raise HTTPException(403)
    v = store.get_version(session_id, version)
    if not v:
        raise HTTPException(404, f"Version {version} not found")
    return {
        "version": version,
        "title": v.get("blog_title", ""),
        "content": v.get("humanized_content", ""),
        "word_count": v.get("blog_word_count", 0),
        "is_current": v.get("is_current", False),
    }


@router.post("/chat/restore-version")
async def restore_version(request: ChatRequest, user_id: str = Depends(get_current_user)):
    sid = request.session_id
    if not sid:
        raise HTTPException(400, "session_id required")
    s = store.get_session(sid)
    if not s:
        raise HTTPException(404, "Session not found")
    if s.user_id != user_id:
        raise HTTPException(403)
    try:
        version = int(request.message)  # message field carries the version number
    except (ValueError, TypeError):
        raise HTTPException(400, "Invalid version number")
    draft = store.restore_version(sid, version)
    if not draft:
        raise HTTPException(404, f"Version {version} not found")
    return {
        "success": True,
        "version": draft.get("current_version"),
        "title": draft.get("blog_title", ""),
        "content": draft.get("humanized_content", ""),
        "word_count": draft.get("blog_word_count", 0),
    }


# ---------------------------------------------------------------------------
# Export to DOCX
# ---------------------------------------------------------------------------


@router.get("/chat/export-docx")
async def export_docx(session_id: str, version: Optional[int] = None, user_id: str = Depends(get_current_user)):
    _get_owned_session(session_id, user_id)
    if version is not None:
        draft = store.get_version(session_id, version)
    else:
        draft = store.get_draft(session_id)
    if not draft or not draft.get("humanized_content"):
        raise HTTPException(400, "No blog content to export")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, "python-docx not installed")

    title = draft.get("blog_title", "Untitled")
    content = draft["humanized_content"]
    keywords = draft.get("keywords", "")
    meta_desc = draft.get("meta_description", "")

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta info
    if meta_desc:
        meta = doc.add_paragraph()
        meta.add_run("Meta: ").bold = True
        meta.add_run(meta_desc)
        meta.style = 'Intense Quote'
    if keywords:
        kw = doc.add_paragraph()
        kw.add_run("Keywords: ").bold = True
        kw.add_run(keywords)
        kw.style = 'Intense Quote'

    doc.add_paragraph("")  # spacer

    # Parse markdown into docx paragraphs
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            # Skip H1 — already used as document title
            pass

        # Bullet lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, text)

        # Numbered lists
        elif len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:5]:
            text = stripped.split(". ", 1)[1]
            p = doc.add_paragraph(style='List Number')
            _add_formatted_runs(p, text)

        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # Blockquote
        elif stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph(style='Intense Quote')
            _add_formatted_runs(p, text)

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)

        i += 1

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()

    safe_name = re.sub(r"[^\w\s-]", "", title.lower())[:50]
    safe_name = re.sub(r"[\s-]+", "-", safe_name)
    filename = f"{safe_name or 'blog'}.docx"

    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _add_formatted_runs(paragraph, text: str):
    """Add bold/italic runs from markdown-ish text to a docx paragraph."""
    import re as _re
    # Split on **bold** and *italic* patterns
    parts = _re.split(r'(\*\*.+?\*\*|\*.+?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Sessions
# ---------------------------------------------------------------------------


@router.get("/sessions", response_model=list[SessionResponse])
async def list_sessions(user_id: str = Depends(get_current_user)):
    return store.list_sessions(user_id)


@router.get("/sessions/{session_id}", response_model=SessionResponse)
async def get_session(session_id: str, user_id: str = Depends(get_current_user)):
    s = store.get_session(session_id)
    if not s:
        raise HTTPException(404, "Not found")
    if s.user_id != user_id:
        raise HTTPException(403)
    return s


@router.delete("/sessions/{session_id}")
async def delete_session(session_id: str, user_id: str = Depends(get_current_user)):
    s = store.get_session(session_id)
    if not s:
        raise HTTPException(404, "Not found")
    if s.user_id != user_id:
        raise HTTPException(403)
    store.delete_session(session_id)
    return {"detail": "Deleted"}
